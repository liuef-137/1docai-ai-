import json
import re
import hashlib
import os
import uuid
import smtplib
from email.message import EmailMessage
from io import BytesIO
import requests as http_requests
from datetime import datetime, date, timedelta
from flask import (
    Blueprint, request, jsonify, current_app, make_response, render_template,
    send_from_directory, send_file
)
from werkzeug.utils import secure_filename
from sqlalchemy import or_, func
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired

from models import db, User, Analysis, RiskPreference, ContractCompare, Feedback, Conversation, UserQuota, GuestQuota, Notification, UserNotification
from auth import create_token, decode_token, get_current_user, admin_required
from multilingual import (
    detect_language,
    build_analysis_prompt,
    build_few_shot_messages,
    get_review_stance_prompt,
    get_plain_language_injection,
)

api_bp = Blueprint('api', __name__)


def _optional_user():
    token = request.cookies.get('token')
    if not token:
        auth_header = request.headers.get('Authorization', '')
        token = auth_header[7:] if auth_header.startswith('Bearer ') else None
    return decode_token(token) if token else None


def _guest_id():
    guest_id = request.cookies.get('docai_guest_id')
    return guest_id or uuid.uuid4().hex


def _client_ip_hash():
    """Hash the client IP; only use forwarded addresses from configured proxies."""
    proxy_hops = max(0, int(current_app.config.get('TRUSTED_PROXY_HOPS', 0) or 0))
    if proxy_hops and len(request.access_route) > proxy_hops:
        ip = request.access_route[-(proxy_hops + 1)]
    else:
        ip = request.remote_addr or 'unknown'
    secret = current_app.config.get('SECRET_KEY', '')
    return hashlib.sha256(f'{secret}:{ip}'.encode('utf-8')).hexdigest()


def _email_verified(user):
    return user is not None and user.email_verified is not False


def _email_token_serializer():
    return URLSafeTimedSerializer(current_app.config['SECRET_KEY'], salt='docai-email-verification')


def _verification_url(user_id):
    token = _email_token_serializer().dumps(str(user_id))
    base_url = current_app.config.get('PUBLIC_BASE_URL') or request.url_root.rstrip('/')
    return f'{base_url}/api/auth/verify-email?token={token}'


def _send_verification_email(user):
    url = _verification_url(user.id)
    host = current_app.config.get('SMTP_HOST')
    if not host:
        current_app.logger.warning('Email verification URL for %s: %s', user.email, url)
        return url

    message = EmailMessage()
    message['Subject'] = 'Verify your DocAI email'
    message['From'] = current_app.config.get('SMTP_FROM') or current_app.config.get('SMTP_USERNAME')
    message['To'] = user.email
    message.set_content(f'Please verify your DocAI email by opening this link:\n\n{url}\n\nThe link expires in 24 hours.')
    with smtplib.SMTP(host, current_app.config.get('SMTP_PORT', 587), timeout=15) as server:
        server.starttls()
        server.login(current_app.config.get('SMTP_USERNAME'), current_app.config.get('SMTP_PASSWORD'))
        server.send_message(message)
    return url


def _get_anonymous_user():
    user = User.query.filter_by(role='anonymous').first()
    if user:
        return user
    user = User(
        username='anonymous',
        email='anonymous@system.docai.local',
        role='anonymous',
        email_verified=True,
    )
    user.set_password(uuid.uuid4().hex)
    db.session.add(user)
    db.session.commit()
    return user


def _text_limit_for_user(current_user):
    if not current_user:
        return current_app.config.get('GUEST_MAX_TEXT_LENGTH', 1000)
    if current_user.role == 'admin':
        return current_app.config.get('MAX_TEXT_LENGTH', 20000)
    plan_limits = {
        'starter': 2000,
        'pro': current_app.config.get('STANDARD_MAX_TEXT_LENGTH', 5000),
        'business': current_app.config.get('MAX_TEXT_LENGTH', 20000),
    }
    return plan_limits.get(current_user.plan or 'free', current_app.config.get('FREE_MAX_TEXT_LENGTH', 1000))


def _maybe_grant_referral_reward(user):
    if not user or not _email_verified(user) or not user.referred_by_user_id or user.referral_reward_granted:
        return
    referrer = db.session.get(User, user.referred_by_user_id)
    if not referrer or referrer.id == user.id:
        return
    bonus = current_app.config.get('REFERRAL_BONUS_CREDITS', 2)
    UserQuota.add_bonus_credits(referrer.id, bonus, commit=False)
    user.referral_reward_granted = True
    db.session.commit()


def _refund_analysis_quota(current_user, guest_id):
    if current_user:
        UserQuota.refund(current_user.id, 'analysis')
    else:
        GuestQuota.refund(guest_id)


def _usage_fields(response):
    usage = getattr(response, 'usage', None) or {}
    return {
        'ai_provider': getattr(response, 'provider', None),
        'prompt_tokens': usage.get('prompt_tokens'),
        'completion_tokens': usage.get('completion_tokens'),
        'total_tokens': usage.get('total_tokens'),
    }

# ---------------------------------------------------------------------------
# DeepSeek API helper
# ---------------------------------------------------------------------------

def call_deepseek(messages, stream=False):
    """
    Call the DeepSeek chat completions API.

    Args:
        messages: list of dicts with 'role' and 'content'
        stream: if True, return a generator yielding content chunks

    Returns:
        str (full content) or generator of str chunks

    Raises:
        RuntimeError on API errors
    """
    config = current_app.config
    providers = []
    relay_key = config.get('RELAY_API_KEY', '')
    relay_base = config.get('RELAY_API_BASE_URL', '')
    if relay_key and relay_base:
        retries = max(1, config.get('RELAY_MAX_RETRIES', 2))
        providers.extend([(
            '中转站', relay_base, relay_key, config.get('RELAY_MODEL', 'deepseek-chat'),
            config.get('RELAY_TIMEOUT_SECONDS', 20),
        )] * retries)

    deepseek_key = config.get('DEEPSEEK_API_KEY', '')
    if deepseek_key:
        providers.append((
            'DeepSeek', config.get('DEEPSEEK_BASE_URL', 'https://api.deepseek.com').rstrip('/'),
            deepseek_key, config.get('DEEPSEEK_MODEL', 'deepseek-chat'),
            config.get('DEEPSEEK_TIMEOUT_SECONDS', 60),
        ))
    if not providers:
        raise RuntimeError('AI API 未配置，请设置中转站或 DeepSeek API')

    errors = []
    payload = {
        'model': None,
        'messages': messages,
        'temperature': 0.3,
        'max_tokens': 4096,
        'stream': stream,
    }
    if stream:
        payload['stream_options'] = {'include_usage': True}

    if stream:
        class AIStream:
            def __init__(self, factory):
                self._factory = factory
                self._iterator = None
                self.provider = None
                self.usage = {}

            def __iter__(self):
                if self._iterator is None:
                    self._iterator = self._factory()
                return self

            def __next__(self):
                if self._iterator is None:
                    self._iterator = self._factory()
                return next(self._iterator)

        def _stream():
            # A relay can accept the connection and then return no usable SSE data.
            # Keep fallback inside the generator so failures during iteration can
            # still move to the next provider before anything reaches the client.
            stream_errors = []
            for provider_name, base_url, api_key, model, timeout in providers:
                resp = None
                emitted = False
                try:
                    url = f'{base_url.rstrip("/")}/chat/completions'
                    headers = {'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'}
                    request_payload = dict(payload, model=model)
                    resp = http_requests.post(
                        url, headers=headers, json=request_payload,
                        timeout=timeout, stream=True,
                    )
                    if resp.status_code != 200:
                        stream_errors.append(f'{provider_name} HTTP {resp.status_code}')
                        continue
                    for line in resp.iter_lines(decode_unicode=True):
                        if not line:
                            continue
                        if not line.startswith('data: '):
                            continue
                        data_str = line[6:]
                        if data_str.strip() == '[DONE]':
                            break
                        try:
                            data = json.loads(data_str)
                            if data.get('usage'):
                                stream_result.usage = data['usage']
                            delta = data.get('choices', [{}])[0].get('delta', {})
                            content = delta.get('content', '')
                            if content:
                                emitted = True
                                yield content
                        except (json.JSONDecodeError, IndexError, KeyError, TypeError):
                            continue
                    if emitted:
                        stream_result.provider = provider_name
                        return
                    stream_errors.append(f'{provider_name}: empty stream')
                except http_requests.RequestException as exc:
                    if emitted:
                        raise RuntimeError(f'{provider_name} 流式响应中断: {exc}') from exc
                    stream_errors.append(f'{provider_name}: {exc}')
                finally:
                    if resp is not None:
                        resp.close()
            raise RuntimeError('AI 服务暂不可用：' + '; '.join(stream_errors[-3:]))
        stream_result = AIStream(_stream)
        return stream_result

    for provider_name, base_url, api_key, model, timeout in providers:
        url = f'{base_url.rstrip("/")}/chat/completions'
        headers = {'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'}
        try:
            candidate = http_requests.post(
                url, headers=headers, json=dict(payload, model=model), timeout=timeout,
            )
            if candidate.status_code != 200:
                errors.append(f'{provider_name} HTTP {candidate.status_code}')
                continue
            try:
                body = candidate.json()
                content = body['choices'][0]['message']['content']
                class AIResponse(str):
                    def __new__(cls, value, provider, usage):
                        result = str.__new__(cls, value)
                        result.provider = provider
                        result.usage = usage or {}
                        return result
                return AIResponse(content, provider_name, body.get('usage'))
            except (KeyError, IndexError, TypeError, ValueError) as exc:
                errors.append(f'{provider_name} invalid response: {exc}')
        except http_requests.RequestException as exc:
            errors.append(f'{provider_name}: {exc}')
    raise RuntimeError('AI 服务暂不可用：' + '; '.join(errors[-3:]))


def _safe_parse_json(text):
    """Try to extract and parse JSON from an AI response text."""
    if not text:
        return None

    # Try direct parse first
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Try to find JSON block in markdown code fences
    # Match ```json ... ``` or ``` ... ```
    patterns = [
        r'```json\s*\n?(.*?)\n?\s*```',
        r'```\s*\n?(.*?)\n?\s*```',
        r'\{[\s\S]*\}',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.DOTALL)
        if match:
            candidate = match.group(1) if match.lastindex else match.group(0)
            try:
                return json.loads(candidate)
            except json.JSONDecodeError:
                continue

    return None


# ---------------------------------------------------------------------------
# Multilingual prompts are now managed in multilingual.py
# ---------------------------------------------------------------------------

COMPARE_PROMPT = """你是一位资深合同律师。请对比分析以下两份合同文本的差异，并从法律风险角度解读这些变更。

原始合同：
{original_text}

修改后合同：
{modified_text}

请以严格的 JSON 格式返回分析结果，不要包含任何其他文字说明。JSON 结构如下：
{
  "changes": [
    {
      "type": "<added/modified/deleted>",
      "location": "<变更位置描述>",
      "original": "<原始内容>",
      "modified": "<修改后内容>",
      "impact": "<变更影响分析>",
      "risk_level": "<high/medium/low/none>",
      "suggestion": "<建议>"
    }
  ],
  "overall_assessment": "<整体评估>",
  "risk_summary": "<风险总结>"
}

请确保返回有效的 JSON。"""

COMPARE_PROMPT_EN = """You are a senior contract lawyer. Compare the following two contract texts and explain the legal risk impact of the changes.

Original contract:
{original_text}

Modified contract:
{modified_text}

Return strictly valid JSON and no extra text. Use this structure:
{
  "changes": [
    {
      "type": "<added/modified/deleted>",
      "location": "<where the change appears>",
      "original": "<original content>",
      "modified": "<modified content>",
      "impact": "<legal impact analysis>",
      "risk_level": "<high/medium/low/none>",
      "suggestion": "<suggestion>"
    }
  ],
  "overall_assessment": "<overall assessment>",
  "risk_summary": "<risk summary>"
}

All values intended for users must be written in English. Ensure the response is valid JSON."""

FOLLOWUP_PROMPT = """你是 DocAI 的合同分析助手。用户之前对一份合同进行了分析，现在想继续追问。

请根据之前的分析结果和用户的追问，给出专业、清晰、通俗易懂的回答。

重要规则：
1. 回答要具体，引用合同中的具体条款和之前的分析结果
2. 如果用户问的是法律建议，请声明"以下仅为 AI 分析参考，不构成法律意见"
3. 用通俗语言解释专业法律概念
4. 如果追问与当前分析无关，请礼貌地引导回合同分析话题
5. 回答控制在 500 字以内，重点突出"""

FOLLOWUP_PROMPT_EN = """You are DocAI's contract analysis assistant. The user previously analyzed a contract and now has a follow-up question.

Answer based on the previous analysis and the user's question. Be professional, clear, and easy to understand.

Rules:
1. Be specific and refer to concrete clauses and prior analysis when relevant.
2. If the user asks for legal advice, state that the answer is AI analysis for reference only and not legal advice.
3. Explain legal concepts in plain English.
4. If the question is unrelated to the current contract analysis, politely guide the user back to contract analysis.
5. Keep the answer within 500 words and focus on the key points."""

# ---------------------------------------------------------------------------
# Contract Type Detection
# ---------------------------------------------------------------------------

CONTRACT_TYPE_PROMPT_ZH = """请识别以下合同文本的类型。只返回一个 JSON，不要有其他文字：
{{"type": "<contract_type>", "confidence": <0.0-1.0>}}

合同类型可选值：
- labor: 劳动合同
- rental: 房屋租赁合同
- purchase: 买卖合同 / 购销合同
- service: 服务合同 / 技术服务合同
- nda: 保密协议 / 保密合同
- loan: 借款合同 / 借贷合同
- partnership: 合伙协议 / 合作协议
- franchise: 加盟合同 / 特许经营合同
- agency: 委托合同 / 代理合同
- construction: 建设工程合同
- insurance: 保险合同
- other: 其他

合同文本：
{contract_text}
"""

CONTRACT_TYPE_PROMPT_EN = """Identify the type of the following contract. Return only JSON, no other text:
{{"type": "<contract_type>", "confidence": <0.0-1.0>}}

Allowed contract types:
- labor: employment agreement / labor contract
- rental: lease / rental agreement
- purchase: sale / purchase contract
- service: service agreement / technical service contract
- nda: non-disclosure agreement / confidentiality agreement
- loan: loan / lending agreement
- partnership: partnership / cooperation agreement
- franchise: franchise agreement
- agency: agency / representation agreement
- construction: construction contract
- insurance: insurance contract
- other: other

Contract text:
{contract_text}
"""

CONTRACT_TYPE_FOCUS = {
    'labor': {
        'zh': '劳动合同专项审查',
        'en': 'Employment Agreement Review',
        'focus': {
            'zh': '重点关注：试用期时长与工资、社保缴纳、竞业限制补偿、加班费计算、解除条件、违约金合法性、经济补偿金标准。注意《劳动合同法》对上述条款的强制性规定。',
            'en': 'Focus on: probation period and pay, social insurance, non-compete compensation, overtime calculation, termination conditions, validity of penalties, and severance standards. Note mandatory rules under employment law.',
        },
    },
    'rental': {
        'zh': '房屋租赁合同专项审查',
        'en': 'Lease / Rental Agreement Review',
        'focus': {
            'zh': '重点关注：租期与违约金、押金退还条件、维修责任划分、转租限制、装修补偿、提前解约权、面积与用途约定。注意《民法典》关于租赁合同的规定。',
            'en': 'Focus on: lease term and penalties, deposit refund conditions, maintenance responsibilities, subletting restrictions, renovation compensation, early termination rights, and area/use clauses.',
        },
    },
    'purchase': {
        'zh': '买卖合同专项审查',
        'en': 'Sale / Purchase Contract Review',
        'focus': {
            'zh': '重点关注：标的物描述、价格与支付方式、交付条件、质量保证期、验收标准、违约责任、所有权转移时机、风险承担。注意《民法典》买卖合同章节。',
            'en': 'Focus on: subject matter description, price and payment, delivery conditions, warranty period, acceptance criteria, breach remedies, title transfer timing, and risk allocation.',
        },
    },
    'service': {
        'zh': '技术服务合同专项审查',
        'en': 'Service Agreement Review',
        'focus': {
            'zh': '重点关注：服务范围与边界、交付里程碑、验收标准、知识产权归属、保密条款、人员配置、付款节点、维护期与SLA、违约与解约。',
            'en': 'Focus on: scope and boundaries, deliverables/milestones, acceptance criteria, IP ownership, confidentiality, staffing, payment milestones, maintenance/SLA, termination and breach.',
        },
    },
    'nda': {
        'zh': '保密协议专项审查',
        'en': 'NDA / Confidentiality Review',
        'focus': {
            'zh': '重点关注：保密范围界定、保密期限（是否过长）、例外条款、违约赔偿计算、竞业限制关联、信息返还/销毁义务、员工离职后的保密义务。',
            'en': 'Focus on: definition of confidential information, confidentiality period (overly long?), exceptions, damages calculation, non-compete linkage, information return/destruction, and post-employment obligations.',
        },
    },
    'loan': {
        'zh': '借款合同专项审查',
        'en': 'Loan Agreement Review',
        'focus': {
            'zh': '重点关注：借款利率（是否超过LPR 4倍/司法保护上限）、担保方式、还款计划、提前还款违约金、逾期罚息、抵押物处置、保证人责任范围。',
            'en': 'Focus on: interest rate (usury limits), security/collateral, repayment schedule, prepayment penalties, default interest, collateral disposal, and guarantor scope.',
        },
    },
    'other': {
        'zh': '合同通用审查',
        'en': 'General Contract Review',
        'focus': {
            'zh': '请按照通用合同审查标准进行分析。',
            'en': 'Analyze according to general contract review standards.',
        },
    },
}


def _robust_json_parse(text):
    """
    Parse JSON from LLM output with resilience.
    Handles: markdown code blocks, trailing commas, comments, BOM.
    """
    if not text:
        return None

    # Remove BOM
    text = text.strip('\ufeff\ufffe').strip()

    # Extract JSON from markdown code block if present
    json_match = re.search(r'```(?:json)?\s*\n?(.*?)\n?\s*```', text, re.DOTALL)
    if json_match:
        text = json_match.group(1).strip()

    # Remove single-line comments
    text = re.sub(r'//.*$', '', text, flags=re.MULTILINE)

    # Remove multi-line comments (/* ... */)
    text = re.sub(r'/\*.*?\*/', '', text, flags=re.DOTALL)

    # Remove trailing commas before } or ]
    text = re.sub(r',\s*([}\]])', r'\1', text)

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Try to find JSON object in the text
    brace_start = text.find('{')
    if brace_start >= 0:
        try:
            return json.loads(text[brace_start:])
        except json.JSONDecodeError:
            pass

    # Last resort: try to fix common issues
    try:
        # Remove control characters except newline
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
        return json.loads(text)
    except json.JSONDecodeError:
        return None


def _detect_contract_type(contract_text, api_key, language='zh'):
    """Detect contract type using a lightweight AI call."""
    truncated = contract_text[:3000] if len(contract_text) > 3000 else contract_text
    prompt_template = CONTRACT_TYPE_PROMPT_ZH if language == 'zh' else CONTRACT_TYPE_PROMPT_EN
    system_msg = (
        '你是一个合同类型分类器。只返回 JSON。'
        if language == 'zh' else
        'You are a contract type classifier. Return JSON only.'
    )
    messages = [
        {'role': 'system', 'content': system_msg},
        {'role': 'user', 'content': prompt_template.format(contract_text=truncated)},
    ]
    try:
        result = call_deepseek(messages, stream=False)
        parsed = _robust_json_parse(result)
        if parsed and 'type' in parsed:
            return parsed.get('type', 'other'), parsed.get('confidence', 0.5)
    except Exception:
        pass
    return 'other', 0.0


def _compute_text_hash(text):
    """Compute a stable hash for contract text to enable caching."""
    return hashlib.sha256(text.encode('utf-8')).hexdigest()[:16]


# ---------------------------------------------------------------------------
# Sensitive Data Desensitization
# ---------------------------------------------------------------------------

SENSITIVE_PATTERNS = [
    # ID card (18 digits, last may be X)
    ('id_card', r'\b([1-9]\d{5})(19|20)\d{2}(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])\d{3}[\dXx]\b'),
    # Mobile phone (China)
    ('phone', r'\b1[3-9]\d{9}\b'),
    # Bank card (16-19 digits)
    ('bank_card', r'\b\d{16,19}\b'),
    # Email (basic)
    ('email', r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'),
]

def _desensitize_text(text):
    """
    Mask sensitive data in contract text for storage/display.
    Returns (masked_text, mask_map) where mask_map describes what was masked.
    """
    mask_map = []
    masked = text

    for field_type, pattern in SENSITIVE_PATTERNS:
        for match in re.finditer(pattern, masked):
            original = match.group(0)
            start, end = match.start(), match.end()

            # Check if this region was already masked (avoid double-masking)
            before = masked[:start]
            if '****' in before[max(0, start-20):start]:
                continue

            if field_type == 'id_card':
                replacement = original[:4] + '****' + original[-4:]
            elif field_type == 'phone':
                replacement = original[:3] + '****' + original[-4:]
            elif field_type == 'bank_card':
                replacement = original[:4] + '****' + original[-4:]
            elif field_type == 'email':
                at_idx = original.index('@')
                name_part = original[:at_idx]
                replacement = name_part[:2] + '****' + original[at_idx:]
            else:
                replacement = '****'

            masked = masked[:start] + replacement + masked[end:]
            mask_map.append({'type': field_type, 'original': original, 'masked': replacement})

    return masked, mask_map


def _find_cached_analysis(user_id, text_hash, mode):
    """Find a recent cached analysis for the same text."""
    # Look in last 24 hours
    since = datetime.utcnow() - timedelta(hours=24)
    cached = Analysis.query.filter(
        Analysis.user_id == user_id,
        Analysis.created_at >= since
    ).order_by(Analysis.created_at.desc()).first()

    if cached and cached.text_hash == text_hash and cached.analysis_mode == mode:
        return cached
    return None


def _get_mode_prompt(mode):
    prompts = {
        'risk': RISK_ANALYSIS_PROMPT,
        'summary': SUMMARY_PROMPT,
        'plain': PLAIN_LANGUAGE_PROMPT,
    }
    return prompts.get(mode, RISK_ANALYSIS_PROMPT)


def _inject_preferences(system_prompt, preferences):
    """Inject user risk preferences into the system prompt."""
    if not preferences:
        return system_prompt
    pref_text = '、'.join(preferences)
    injection = f"\n\n用户设定的风控红线: {pref_text}。请特别关注这些条款是否触碰红线。"
    return system_prompt + injection


# ---------------------------------------------------------------------------
# Auth Routes
# ---------------------------------------------------------------------------

@api_bp.route('/api/auth/register', methods=['POST'])
def register():
    if not GuestQuota.check_rate_limit(_client_ip_hash(), current_app.config.get('GUEST_RATE_LIMIT_PER_MINUTE', 5)):
        return jsonify({'error': '请求过于频繁，请稍后再试', 'rate_limited': True}), 429
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or '').strip()
    email = (data.get('email') or '').strip()
    password = data.get('password') or ''
    referral_code = (data.get('referral_code') or '').strip().upper()

    if not username or not email or not password:
        return jsonify({'error': '用户名、邮箱和密码不能为空'}), 400

    if not re.match(r'^[^\s@]+@[^\s@]+\.[^\s@]+$', email):
        return jsonify({'error': '请输入有效的邮箱地址'}), 400

    if len(username) < 2 or len(username) > 80:
        return jsonify({'error': '用户名长度需在 2-80 个字符之间'}), 400

    if len(password) < 6:
        return jsonify({'error': '密码长度不能少于 6 个字符'}), 400

    if User.query.filter_by(username=username).first():
        return jsonify({'error': '用户名已存在'}), 409

    if User.query.filter_by(email=email).first():
        return jsonify({'error': '邮箱已被注册'}), 409

    user = User(username=username, email=email, email_verified=False)
    user.set_password(password)

    try:
        db.session.add(user)
        db.session.commit()
        invite_prefix = current_app.config.get('REFERRAL_CODE_PREFIX', 'DCAI')
        user.invite_code = f'{invite_prefix}{user.id:06d}'
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'注册失败: {str(e)}'}), 500

    if referral_code:
        referrer = User.query.filter(func.upper(User.invite_code) == referral_code).first()
        if referrer and referrer.id != user.id:
            try:
                user.referred_by_user_id = referrer.id
                db.session.add(user)
                db.session.commit()
            except Exception:
                db.session.rollback()
                current_app.logger.exception('Failed to apply referral bonus for user %s', user.id)

    try:
        _send_verification_email(user)
    except Exception:
        current_app.logger.exception('Failed to send verification email for user %s', user.id)

    token = create_token(user.id)
    resp = make_response(jsonify({
        'message': '注册成功',
        'token': token,
        'user': user.to_dict(),
        'email_verification_required': True,
    }))
    resp.set_cookie('token', token, httponly=True, max_age=60 * 60 * 24 * 7, samesite='Lax')
    return resp, 201


@api_bp.route('/api/auth/login', methods=['POST'])
def login():
    data = request.get_json(silent=True) or {}
    account = (data.get('username') or '').strip()
    password = data.get('password') or ''

    if not account or not password:
        return jsonify({'error': '用户名或邮箱和密码不能为空'}), 400

    user = User.query.filter(or_(User.username == account, User.email == account)).first()
    if not user or not user.check_password(password):
        return jsonify({'error': '用户名、邮箱或密码错误'}), 401

    token = create_token(user.id)
    if _email_verified(user):
        try:
            UserQuota.ensure_daily_login_bonus(user.id)
        except Exception:
            current_app.logger.exception('Failed to grant login bonus for user %s', user.id)
    resp = make_response(jsonify({
        'message': '登录成功',
        'token': token,
        'user': user.to_dict(),
        'email_verification_required': not _email_verified(user),
    }))
    resp.set_cookie('token', token, httponly=True, max_age=60 * 60 * 24 * 7, samesite='Lax')
    return resp


@api_bp.route('/api/auth/verify-email', methods=['GET'])
def verify_email():
    token = request.args.get('token', '')
    try:
        user_id = int(_email_token_serializer().loads(token, max_age=60 * 60 * 24))
    except (BadSignature, SignatureExpired, ValueError, TypeError):
        return jsonify({'error': '邮箱验证链接无效或已过期'}), 400

    user = db.session.get(User, user_id)
    if not user:
        return jsonify({'error': '用户不存在'}), 404
    user.email_verified = True
    db.session.commit()
    UserQuota.ensure_daily_login_bonus(user.id)
    return jsonify({'message': '邮箱验证成功，现在可以使用完整功能。'})


@api_bp.route('/api/auth/resend-verification', methods=['POST'])
@get_current_user()
def resend_verification(current_user):
    if _email_verified(current_user):
        return jsonify({'message': '邮箱已经验证'}), 200
    try:
        _send_verification_email(current_user)
    except Exception:
        current_app.logger.exception('Failed to resend verification email for user %s', current_user.id)
        return jsonify({'error': '验证邮件发送失败，请稍后重试'}), 502
    return jsonify({'message': '验证邮件已发送'})


@api_bp.route('/api/auth/me', methods=['GET'])
@get_current_user()
def get_me(current_user):
    return jsonify({'user': current_user.to_dict()})


@api_bp.route('/api/auth/logout', methods=['POST'])
def logout():
    resp = make_response(jsonify({'message': '已退出登录'}))
    resp.delete_cookie('token')
    return resp


# 头像存储目录
AVATAR_FOLDER = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'assets', 'avatars'))
AVATAR_ALLOWED = {'jpg', 'jpeg', 'png', 'webp'}
AVATAR_MAX_SIZE = 2 * 1024 * 1024  # 2MB


@api_bp.route('/api/auth/avatar', methods=['POST'])
@get_current_user()
def upload_avatar(current_user):
    """上传/更新当前用户头像。"""
    if 'avatar' not in request.files:
        return jsonify({'error': '未选择文件'}), 400
    file = request.files['avatar']
    if not file or not file.filename:
        return jsonify({'error': '未选择文件'}), 400
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in AVATAR_ALLOWED:
        return jsonify({'error': '仅支持 JPG / PNG / WEBP 格式'}), 400
    # 读取并校验大小
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > AVATAR_MAX_SIZE:
        return jsonify({'error': '图片大小不能超过 2MB'}), 400
    # 确保目录存在
    os.makedirs(AVATAR_FOLDER, exist_ok=True)
    # 删除旧头像文件（仅限本目录下的）
    if current_user.avatar:
        old_name = os.path.basename(current_user.avatar)
        if old_name.startswith('avatar_'):
            try:
                old_path = os.path.join(AVATAR_FOLDER, old_name)
                if os.path.exists(old_path):
                    os.remove(old_path)
            except Exception:
                pass
    # 保存新头像
    filename = f'avatar_{current_user.id}_{uuid.uuid4().hex[:8]}.{ext}'
    save_path = os.path.join(AVATAR_FOLDER, filename)
    file.save(save_path)
    # 更新数据库（存相对 URL）
    avatar_url = f'/assets/avatars/{filename}'
    current_user.avatar = avatar_url
    db.session.commit()
    return jsonify({'message': '头像更新成功', 'avatar': avatar_url, 'user': current_user.to_dict()})


@api_bp.route('/api/auth/avatar', methods=['DELETE'])
@get_current_user()
def delete_avatar(current_user):
    """删除当前用户头像，回退为首字母占位。"""
    if current_user.avatar:
        old_name = os.path.basename(current_user.avatar)
        if old_name.startswith('avatar_'):
            try:
                old_path = os.path.join(AVATAR_FOLDER, old_name)
                if os.path.exists(old_path):
                    os.remove(old_path)
            except Exception:
                pass
        current_user.avatar = None
        db.session.commit()
    return jsonify({'message': '头像已移除', 'user': current_user.to_dict()})


# ---------------------------------------------------------------------------
# Analysis Routes
# ---------------------------------------------------------------------------

@api_bp.route('/api/analysis', methods=['POST'])
def create_analysis():
    current_user = _optional_user()
    guest_id = None if current_user else _guest_id()
    ip_hash = _client_ip_hash()
    data = request.get_json(silent=True) or {}
    text = (data.get('text') or '').strip()
    mode = (data.get('mode') or 'risk').strip()
    filename = (data.get('filename') or '').strip() or None
    review_stance = (data.get('review_stance') or '').strip()  # 'party_a' or 'party_b'

    if not text:
        return jsonify({'error': '合同文本不能为空'}), 400

    if review_stance and review_stance not in ('party_a', 'party_b'):
        review_stance = ''

    if current_user and current_user.role != 'admin' and not _email_verified(current_user):
        return jsonify({
            'error': '请先验证邮箱后再开始分析',
            'email_verification_required': True,
        }), 403

    max_len = _text_limit_for_user(current_user)
    if len(text) > max_len:
        return jsonify({
            'error': f'合同文本过长，请控制在 {max_len} 字以内',
            'count': len(text),
            'max': max_len,
        }), 400

    if mode not in ('risk', 'summary', 'plain'):
        mode = 'risk'

    # Compute text hash for caching
    text_hash = _compute_text_hash(text)

    # Check cache — if found, return cached result immediately
    cached = _find_cached_analysis(current_user.id, text_hash, mode) if current_user else None
    if cached:
        return jsonify({
            'message': '分析完成（缓存）',
            'analysis': cached.to_dict(),
            'cached': True,
        }), 200

    if current_user:
        allowed, _, _ = UserQuota.check_and_increment(current_user.id, 'analysis')
    else:
        if not GuestQuota.check_rate_limit(
            ip_hash,
            current_app.config.get('GUEST_RATE_LIMIT_PER_MINUTE', 5),
        ):
            return jsonify({'error': '请求过于频繁，请稍后再试', 'rate_limited': True}), 429
        allowed = GuestQuota.check_and_increment(guest_id, ip_hash)
    if not allowed:
        if current_user:
            return jsonify({
                'error': '今日分析额度已用完，可邀请好友获得额外额度，或明天再来使用。',
                'quota_exhausted': True,
                'require_login': False,
            }), 429

        response = jsonify({
            'error': '今日免费分析次数已用完，请登录后继续使用',
            'require_login': True,
        })
        response.set_cookie('docai_guest_id', guest_id, max_age=60 * 60 * 24 * 365, samesite='Lax')
        return response, 429
    # Detect contract language
    language = detect_language(text)

    # Detect contract type (only for risk mode)
    contract_type = 'other'
    if mode == 'risk':
        api_key = current_app.config.get('DEEPSEEK_API_KEY', '')
        contract_type, _ = _detect_contract_type(text, api_key, language=language)

    # Load user risk preferences for risk analysis mode
    user_prefs = []
    if mode == 'risk' and current_user:
        pref_obj = RiskPreference.query.filter_by(user_id=current_user.id).first()
        if pref_obj:
            user_prefs = pref_obj.get_preferences_list()

    # Build prompt with type-specific focus for risk mode
    if mode == 'risk':
        type_info = CONTRACT_TYPE_FOCUS.get(contract_type, CONTRACT_TYPE_FOCUS['other'])
        type_focus_str = f"\n{type_info.get(language, type_info['zh'])}：{type_info['focus'].get(language, type_info['focus']['zh'])}"
        stance_prompt = get_review_stance_prompt(language, review_stance)
        plain_lang_prompt = get_plain_language_injection(language)
        system_prompt = build_analysis_prompt(
            language, mode,
            type_focus=type_focus_str,
            stance=stance_prompt,
            plain=plain_lang_prompt,
        )
        if user_prefs:
            system_prompt = _inject_preferences(system_prompt, user_prefs)
    else:
        system_prompt = build_analysis_prompt(language, mode)

    messages = [{'role': 'system', 'content': system_prompt}]
    messages.extend(build_few_shot_messages(language, mode))
    messages.append({'role': 'user', 'content': text})

    try:
        ai_response = call_deepseek(messages)
        if not str(ai_response or '').strip():
            raise RuntimeError('AI 返回了空结果，请稍后重试')
    except Exception as e:
        db.session.rollback()
        _refund_analysis_quota(current_user, guest_id)
        return jsonify({'error': str(e)}), 502

    # Parse the AI response
    parsed = _robust_json_parse(ai_response)

    # Extract fields
    score = None
    risk_level = None
    one_line_summary = None
    suggestions_json = None
    risk_items_json = None
    result_json = None

    if parsed and isinstance(parsed, dict):
        score = parsed.get('score')
        risk_level = parsed.get('risk_level')
        one_line_summary = parsed.get('one_line_summary')

        suggestions = parsed.get('suggestions')
        if suggestions is not None:
            try:
                suggestions_json = json.dumps(suggestions, ensure_ascii=False)
            except (TypeError, ValueError):
                pass

        risk_items = parsed.get('risk_items')
        if risk_items is not None:
            try:
                risk_items_json = json.dumps(risk_items, ensure_ascii=False)
            except (TypeError, ValueError):
                pass

        try:
            result_json = json.dumps(parsed, ensure_ascii=False)
        except (TypeError, ValueError):
            result_json = ai_response

        # Fallback: if score/risk_level missing in parsed data
        if score is None and risk_level is None:
            risk_level = 'medium'
    else:
        # JSON parse failed, store raw response
        result_json = ai_response
        one_line_summary = ai_response[:200] if ai_response else None

    # Convert score to int if possible
    if score is not None:
        try:
            score = int(score)
        except (ValueError, TypeError):
            score = None

    # Desensitize contract text before storage
    masked_text, _ = _desensitize_text(text)

    # Create analysis record
    analysis = Analysis(
        user_id=current_user.id if current_user else _get_anonymous_user().id,
        filename=filename,
        contract_text=masked_text,
        text_hash=text_hash,
        contract_type=contract_type,
        language=language,
        analysis_mode=mode,
        result=result_json,
        score=score,
        risk_level=risk_level,
        one_line_summary=one_line_summary,
        suggestions=suggestions_json,
        risk_items=risk_items_json,
        is_anonymous=current_user is None,
        source_ip_hash=ip_hash if current_user is None else None,
        **_usage_fields(ai_response),
    )

    try:
        if current_user:
            db.session.add(analysis)
            db.session.commit()
        else:
            db.session.add(analysis)
            db.session.commit()
    except Exception as e:
        db.session.rollback()
        _refund_analysis_quota(current_user, guest_id)
        return jsonify({'error': f'保存分析结果失败: {str(e)}'}), 500

    if current_user:
        try:
            _maybe_grant_referral_reward(current_user)
        except Exception:
            # Referral accounting must not turn a successful analysis into a
            # failed request or refund an already persisted analysis.
            db.session.rollback()
            current_app.logger.exception('Failed to grant referral reward for user %s', current_user.id)

    response = jsonify({
        'message': '分析完成',
        'analysis': analysis.to_dict(),
    }), 201
    if not current_user: response[0].set_cookie('docai_guest_id', guest_id, max_age=60 * 60 * 24 * 365, samesite='Lax')
    return response


@api_bp.route('/api/analysis/history', methods=['GET'])
@get_current_user()
def analysis_history(current_user):
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    risk_level = request.args.get('risk_level', '').strip()
    mode = request.args.get('mode', '').strip()
    search = request.args.get('search', '').strip()

    per_page = min(max(per_page, 1), 100)

    query = Analysis.query.filter_by(user_id=current_user.id)

    if risk_level and risk_level in ('high', 'medium', 'low'):
        query = query.filter(Analysis.risk_level == risk_level)

    if mode and mode in ('risk', 'summary', 'plain'):
        query = query.filter(Analysis.analysis_mode == mode)

    if search:
        query = query.filter(or_(
            Analysis.filename.contains(search),
            Analysis.contract_text.contains(search),
            Analysis.one_line_summary.contains(search),
        ))

    query = query.order_by(Analysis.created_at.desc())
    pagination = query.paginate(page=page, per_page=per_page, error_out=False)

    items = [a.to_dict() for a in pagination.items]
    for item in items:
        item['display_user'] = '匿名' if item.get('is_anonymous') else item.get('user_id')

    return jsonify({
        'items': items,
        'total': pagination.total,
        'page': pagination.page,
        'per_page': pagination.per_page,
        'pages': pagination.pages,
    })


@api_bp.route('/api/analysis/<int:analysis_id>', methods=['GET'])
@get_current_user()
def get_analysis(current_user, analysis_id):
    analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first()
    if not analysis:
        return jsonify({'error': '分析记录不存在'}), 404
    return jsonify({'analysis': analysis.to_dict()})


@api_bp.route('/api/analysis/<int:analysis_id>', methods=['DELETE'])
@get_current_user()
def delete_analysis(current_user, analysis_id):
    analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first()
    if not analysis:
        return jsonify({'error': '分析记录不存在'}), 404

    try:
        db.session.delete(analysis)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'删除失败: {str(e)}'}), 500

    return jsonify({'message': '已删除'})


@api_bp.route('/api/analysis/<int:analysis_id>/favorite', methods=['POST'])
@get_current_user()
def toggle_favorite(current_user, analysis_id):
    analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first()
    if not analysis:
        return jsonify({'error': '分析记录不存在'}), 404

    analysis.is_favorited = not analysis.is_favorited

    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'操作失败: {str(e)}'}), 500

    return jsonify({
        'message': '已收藏' if analysis.is_favorited else '已取消收藏',
        'is_favorited': analysis.is_favorited,
    })


# ---------------------------------------------------------------------------
# Risk Preference Routes
# ---------------------------------------------------------------------------

@api_bp.route('/api/preferences', methods=['GET'])
@get_current_user()
def get_preferences(current_user):
    pref = RiskPreference.query.filter_by(user_id=current_user.id).first()
    if not pref:
        return jsonify({'preferences': []})
    return jsonify({'preferences': pref.get_preferences_list()})


@api_bp.route('/api/preferences', methods=['PUT'])
@get_current_user()
def update_preferences(current_user):
    data = request.get_json(silent=True) or {}
    preferences = data.get('preferences')

    if preferences is not None and not isinstance(preferences, list):
        return jsonify({'error': 'preferences 必须是字符串数组'}), 400

    pref = RiskPreference.query.filter_by(user_id=current_user.id).first()
    if not pref:
        pref = RiskPreference(user_id=current_user.id)
        db.session.add(pref)

    if preferences is not None:
        pref.set_preferences_list(preferences)

    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'保存偏好设置失败: {str(e)}'}), 500

    return jsonify({
        'message': '偏好设置已更新',
        'preferences': pref.get_preferences_list(),
    })


# ---------------------------------------------------------------------------
# Compare Routes
# ---------------------------------------------------------------------------

@api_bp.route('/api/compare', methods=['POST'])
@get_current_user()
def create_compare(current_user):
    data = request.get_json(silent=True) or {}
    original_text = (data.get('original_text') or '').strip()
    modified_text = (data.get('modified_text') or '').strip()
    language = data.get('language') or detect_language(original_text + '\n' + modified_text)
    if language not in ('zh', 'en'):
        language = 'zh'

    if not original_text or not modified_text:
        return jsonify({'error': '原始合同和修改后合同文本都不能为空'}), 400

    max_len = _text_limit_for_user(current_user)
    if len(original_text) > max_len or len(modified_text) > max_len:
        return jsonify({
            'error': f'合同文本过长，请控制在 {max_len} 字以内',
        }), 400

    allowed, remaining, limit = UserQuota.check_and_increment(current_user.id, 'compare')
    if not allowed:
        period_label = '本月' if current_user.plan in UserQuota.PAID_LIMITS else '可用'
        return jsonify({'error': f'{period_label}对比次数已达上限（{limit}次），请稍后再试', 'remaining': 0, 'daily_limit': limit}), 429

    # Call DeepSeek for comparison analysis
    prompt_template = COMPARE_PROMPT_EN if language == 'en' else COMPARE_PROMPT
    system_prompt = prompt_template.replace('{original_text}', original_text).replace('{modified_text}', modified_text)
    messages = [
        {'role': 'system', 'content': system_prompt},
        {'role': 'user', 'content': 'Compare the differences between the two contracts above.' if language == 'en' else '请对比分析以上两份合同的差异。'},
    ]

    try:
        ai_response = call_deepseek(messages)
        if not str(ai_response or '').strip():
            raise RuntimeError('AI 返回了空结果，请稍后重试')
    except Exception as e:
        db.session.rollback()
        UserQuota.refund(current_user.id, 'compare')
        return jsonify({'error': str(e)}), 502

    parsed = _safe_parse_json(ai_response)
    diff_json = None
    interpretation = None

    if parsed and isinstance(parsed, dict):
        try:
            diff_json = json.dumps(parsed, ensure_ascii=False)
        except (TypeError, ValueError):
            pass
        interpretation = parsed.get('overall_assessment') or parsed.get('risk_summary')
    else:
        interpretation = ai_response

    compare = ContractCompare(
        user_id=current_user.id,
        original_text=original_text,
        modified_text=modified_text,
        diff_result=diff_json,
        ai_interpretation=interpretation,
    )

    try:
        db.session.add(compare)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        UserQuota.refund(current_user.id, 'compare')
        return jsonify({'error': f'保存对比结果失败: {str(e)}'}), 500

    return jsonify({
        'message': '对比分析完成',
        'compare': compare.to_dict(),
    }), 201


@api_bp.route('/api/compare/history', methods=['GET'])
@get_current_user()
def compare_history(current_user):
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    per_page = min(max(per_page, 1), 100)

    query = ContractCompare.query.filter_by(user_id=current_user.id).order_by(
        ContractCompare.created_at.desc()
    )
    pagination = query.paginate(page=page, per_page=per_page, error_out=False)

    items = [c.to_dict() for c in pagination.items]

    return jsonify({
        'items': items,
        'total': pagination.total,
        'page': pagination.page,
        'per_page': pagination.per_page,
        'pages': pagination.pages,
    })


# ---------------------------------------------------------------------------
# Summary Card Route
# ---------------------------------------------------------------------------

@api_bp.route('/api/card/<int:analysis_id>', methods=['GET'])
@get_current_user()
def get_card(current_user, analysis_id):
    analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first()
    if not analysis:
        return jsonify({'error': '分析记录不存在'}), 404

    lang = request.args.get('lang') or getattr(analysis, 'language', None) or 'zh'
    if lang not in ('zh', 'en'):
        lang = 'zh'

    # Render a summary card as HTML
    risk_level_map = {
        'zh': {'high': '高风险', 'medium': '中风险', 'low': '低风险', 'unknown': '未知'},
        'en': {'high': 'High Risk', 'medium': 'Medium Risk', 'low': 'Low Risk', 'unknown': 'Unknown'},
    }[lang]
    risk_level_label = risk_level_map.get(analysis.risk_level, risk_level_map['unknown'])
    mode_map = {
        'zh': {'risk': '风险分析', 'summary': '合同摘要', 'plain': '通俗解读'},
        'en': {'risk': 'Risk Analysis', 'summary': 'Summary', 'plain': 'Plain Language'},
    }[lang]
    mode_label = mode_map.get(analysis.analysis_mode, analysis.analysis_mode)

    card_html = render_template('card.html', analysis=analysis,
                                 risk_level_label=risk_level_label,
                                 mode_label=mode_label,
                                 lang=lang)
    return card_html, 200, {'Content-Type': 'text/html; charset=utf-8'}


# ---------------------------------------------------------------------------
# Streaming analysis endpoint
# ---------------------------------------------------------------------------

@api_bp.route('/api/analysis/stream', methods=['POST'])
@get_current_user()
def stream_analysis(current_user):
    """
    Streaming version of analysis. Returns a SSE (Server-Sent Events) stream.
    """
    data = request.get_json(silent=True) or {}
    text = (data.get('text') or '').strip()
    mode = (data.get('mode') or 'risk').strip()
    filename = (data.get('filename') or '').strip() or None

    if not text:
        return jsonify({'error': '合同文本不能为空'}), 400

    if current_user.role != 'admin' and not _email_verified(current_user):
        return jsonify({'error': '请先验证邮箱后再开始分析', 'email_verification_required': True}), 403

    max_len = _text_limit_for_user(current_user)
    if len(text) > max_len:
        return jsonify({
            'error': f'合同文本过长，请控制在 {max_len} 字以内',
            'count': len(text),
            'max': max_len,
        }), 400

    if mode not in ('risk', 'summary', 'plain'):
        mode = 'risk'

    # Compute text hash for caching
    text_hash = _compute_text_hash(text)

    # Check cache — if found, return cached result immediately
    cached = _find_cached_analysis(current_user.id, text_hash, mode)
    if cached:
        return jsonify({
            'message': '分析完成（缓存）',
            'analysis': cached.to_dict(),
            'cached': True,
        }), 200

    allowed, remaining, limit = UserQuota.check_and_increment(current_user.id, 'analysis')
    if not allowed:
        return jsonify({'error': f'今日分析次数已达上限（{limit}次），请明天再试', 'remaining': 0, 'daily_limit': limit}), 429

    # Detect contract language
    language = detect_language(text)

    # Detect contract type (only for risk mode)
    contract_type = 'other'
    if mode == 'risk':
        api_key = current_app.config.get('DEEPSEEK_API_KEY', '')
        contract_type, _ = _detect_contract_type(text, api_key, language=language)

    # Build prompt with type-specific focus for risk mode
    if mode == 'risk':
        type_info = CONTRACT_TYPE_FOCUS.get(contract_type, CONTRACT_TYPE_FOCUS['other'])
        type_focus_str = f"\n{type_info.get(language, type_info['zh'])}：{type_info['focus'].get(language, type_info['focus']['zh'])}"
        stance_prompt = get_review_stance_prompt(language, '')
        plain_lang_prompt = get_plain_language_injection(language)
        system_prompt = build_analysis_prompt(
            language, mode,
            type_focus=type_focus_str,
            stance=stance_prompt,
            plain=plain_lang_prompt,
        )
    else:
        system_prompt = build_analysis_prompt(language, mode)

    messages = [{'role': 'system', 'content': system_prompt}]
    messages.extend(build_few_shot_messages(language, mode))
    messages.append({'role': 'user', 'content': text})

    def generate():
        full_response = []
        try:
            stream = call_deepseek(messages, stream=True)
            for chunk in stream:
                full_response.append(chunk)
                yield f'data: {json.dumps({"content": chunk}, ensure_ascii=False)}\n\n'

            # After streaming completes, parse and save the analysis
            raw_text = ''.join(full_response)
            parsed = _robust_json_parse(raw_text)

            # Extract fields
            score = None
            risk_level = None
            one_line_summary = None
            suggestions_json = None
            risk_items_json = None
            result_json = None

            if parsed and isinstance(parsed, dict):
                score = parsed.get('score')
                risk_level = parsed.get('risk_level')
                one_line_summary = parsed.get('one_line_summary')

                suggestions = parsed.get('suggestions')
                if suggestions is not None:
                    try:
                        suggestions_json = json.dumps(suggestions, ensure_ascii=False)
                    except (TypeError, ValueError):
                        pass

                risk_items = parsed.get('risk_items')
                if risk_items is not None:
                    try:
                        risk_items_json = json.dumps(risk_items, ensure_ascii=False)
                    except (TypeError, ValueError):
                        pass

                try:
                    result_json = json.dumps(parsed, ensure_ascii=False)
                except (TypeError, ValueError):
                    result_json = raw_text

                if score is None and risk_level is None:
                    risk_level = 'medium'
            else:
                result_json = raw_text
                one_line_summary = raw_text[:200] if raw_text else None

            if score is not None:
                try:
                    score = int(score)
                except (ValueError, TypeError):
                    score = None

            # Save analysis record
            analysis = Analysis(
                user_id=current_user.id,
                filename=filename,
                contract_text=_desensitize_text(text)[0],
                text_hash=text_hash,
                contract_type=contract_type,
                language=language,
                analysis_mode=mode,
                result=result_json,
                score=score,
                risk_level=risk_level,
                one_line_summary=one_line_summary,
                suggestions=suggestions_json,
                risk_items=risk_items_json,
                **_usage_fields(stream),
            )
            db.session.add(analysis)
            db.session.commit()
            try:
                _maybe_grant_referral_reward(current_user)
            except Exception:
                db.session.rollback()
                current_app.logger.exception('Failed to grant referral reward for user %s', current_user.id)

            # Send final event with analysis id
            yield f'data: {json.dumps({"done": True, "analysis_id": analysis.id}, ensure_ascii=False)}\n\n'
            yield 'data: [DONE]\n\n'
        except Exception as e:
            db.session.rollback()
            UserQuota.refund(current_user.id, 'analysis')
            yield f'data: {json.dumps({"error": str(e)}, ensure_ascii=False)}\n\n'

    return current_app.response_class(
        generate(),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
        },
    )


# ---------------------------------------------------------------------------
@api_bp.route('/api/analysis/<int:analysis_id>/conversation', methods=['GET'])
@get_current_user()
def get_conversation(current_user, analysis_id):
    """Get the followup conversation history for an analysis."""
    analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first()
    if not analysis:
        return jsonify({'error': '分析记录不存在'}), 404

    conv = Conversation.query.filter_by(analysis_id=analysis_id).first()
    if not conv:
        return jsonify({'messages': [], 'conversation_id': None})

    return jsonify({'messages': conv.get_messages(), 'conversation_id': conv.id})


# ---------------------------------------------------------------------------
@api_bp.route('/api/analysis/<int:analysis_id>/followup', methods=['POST'])
@get_current_user()
def analysis_followup(current_user, analysis_id):
    """
    Follow-up question on a previous analysis. Supports streaming.
    """
    analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first()
    if not analysis:
        return jsonify({'error': '分析记录不存在'}), 404

    data = request.get_json(silent=True) or {}
    question = (data.get('question') or '').strip()
    language = data.get('language') or analysis.language or 'zh'
    if language not in ('zh', 'en'):
        language = 'zh'
    if not question:
        return jsonify({'error': '请输入追问内容'}), 400

    allowed, remaining, limit = UserQuota.check_and_increment(current_user.id, 'followup')
    if not allowed:
        period_label = '本月' if current_user.plan in UserQuota.PAID_LIMITS else '免费'
        return jsonify({'error': f'{period_label}追问次数已达上限（{limit}次），请稍后再试', 'remaining': 0, 'daily_limit': limit}), 429

    # Get or create conversation for persistence
    conv = Conversation.query.filter_by(analysis_id=analysis_id).first()
    if not conv:
        conv = Conversation(user_id=current_user.id, analysis_id=analysis_id)
        db.session.add(conv)

    # Save the user's question to conversation
    conv.add_message('user', question)

    # Build context from previous analysis
    if language == 'en':
        context_parts = [
            f"Analysis mode: {analysis.analysis_mode}",
            f"Score: {analysis.score}" if analysis.score else None,
            f"Risk level: {analysis.risk_level}" if analysis.risk_level else None,
            f"One-line summary: {analysis.one_line_summary}" if analysis.one_line_summary else None,
        ]
    else:
        context_parts = [
            f"分析模式: {analysis.analysis_mode}",
            f"评分: {analysis.score}" if analysis.score else None,
            f"风险等级: {analysis.risk_level}" if analysis.risk_level else None,
            f"一句话总结: {analysis.one_line_summary}" if analysis.one_line_summary else None,
        ]
    if analysis.risk_items:
        try:
            risk_items = json.loads(analysis.risk_items)
            context_parts.append("Risk clauses:" if language == 'en' else "风险条款:")
            for item in risk_items[:5]:
                if isinstance(item, dict):
                    context_parts.append(f"  - {item.get('clause_name', '')}: {item.get('description', '')}")
        except (json.JSONDecodeError, TypeError):
            pass
    if analysis.suggestions:
        try:
            suggestions = json.loads(analysis.suggestions)
            if suggestions:
                sep = ', ' if language == 'en' else '、'
                context_parts.append(("Suggestions: " if language == 'en' else "改进建议: ") + sep.join(str(s) for s in suggestions[:5]))
        except (json.JSONDecodeError, TypeError):
            pass

    context = '\n'.join(p for p in context_parts if p)

    # Truncate contract text to save tokens
    contract_text = analysis.contract_text or ''
    if len(contract_text) > 3000:
        contract_text = contract_text[:3000] + ('...(truncated)' if language == 'en' else '...(已截断)')

    messages = [
        {'role': 'system', 'content': FOLLOWUP_PROMPT_EN if language == 'en' else FOLLOWUP_PROMPT},
        {'role': 'assistant', 'content': (f'Here is the previous contract analysis:\n\n{context}\n\nOriginal contract text:\n{contract_text}' if language == 'en' else f'以下是之前的合同分析结果：\n\n{context}\n\n原始合同文本：\n{contract_text}')},
        {'role': 'user', 'content': question},
    ]

    full_response = []

    def generate():
        try:
            stream = call_deepseek(messages, stream=True)
            for chunk in stream:
                full_response.append(chunk)
                yield f'data: {json.dumps({"content": chunk}, ensure_ascii=False)}\n\n'
            # Save full response to conversation after streaming completes
            conv.add_message('assistant', ''.join(full_response))
            db.session.commit()
            yield 'data: [DONE]\n\n'
        except Exception as e:
            db.session.rollback()
            UserQuota.refund(current_user.id, 'followup')
            yield f'data: {json.dumps({"error": str(e)}, ensure_ascii=False)}\n\n'

    return current_app.response_class(
        generate(),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
        },
    )


# Admin Routes
# ---------------------------------------------------------------------------

@api_bp.route('/api/admin/stats', methods=['GET'])
@get_current_user()
@admin_required
def admin_stats(current_user):
    total_users = User.query.filter(User.role != 'anonymous').count()
    total_analyses = Analysis.query.count()
    total_notifications = Notification.query.count()
    token_totals = db.session.query(
        func.coalesce(func.sum(Analysis.prompt_tokens), 0),
        func.coalesce(func.sum(Analysis.completion_tokens), 0),
        func.coalesce(func.sum(Analysis.total_tokens), 0),
    ).one()
    anonymous_token_total = db.session.query(
        func.coalesce(func.sum(Analysis.total_tokens), 0)
    ).filter(Analysis.is_anonymous.is_(True)).scalar()

    today_start = datetime.combine(date.today(), datetime.min.time())
    today_analyses = Analysis.query.filter(Analysis.created_at >= today_start).count()

    # Risk distribution
    risk_dist = db.session.query(
        Analysis.risk_level, func.count(Analysis.id)
    ).group_by(Analysis.risk_level).all()
    risk_distribution = {level: count for level, count in risk_dist}

    # ---- Feedback statistics ----
    total_feedback = Feedback.query.count()
    avg_rating = db.session.query(func.avg(Feedback.rating)).filter(Feedback.rating.isnot(None)).scalar()
    avg_rating = round(float(avg_rating), 1) if avg_rating else 0.0

    # NPS: promoters(4-5) - detractors(1-2) / total
    rated = Feedback.query.filter(Feedback.rating.isnot(None)).all()
    if rated:
        promoters = sum(1 for f in rated if f.rating >= 4)
        detractors = sum(1 for f in rated if f.rating <= 2)
        nps = round((promoters - detractors) / len(rated) * 100)
    else:
        nps = 0

    # 7-day daily feedback trend
    daily_feedback = []
    for i in range(6, -1, -1):
        d = date.today() - timedelta(days=i)
        count = Feedback.query.filter(
            func.date(Feedback.created_at) == d
        ).count()
        daily_feedback.append({'date': d.isoformat(), 'count': count})

    # 7-day daily analysis trend
    daily_analysis = []
    for i in range(6, -1, -1):
        d = date.today() - timedelta(days=i)
        count = Analysis.query.filter(
            func.date(Analysis.created_at) == d
        ).count()
        daily_analysis.append({'date': d.isoformat(), 'count': count})

    # Feedback type distribution
    type_dist = db.session.query(
        Feedback.feedback_type, func.count(Feedback.id)
    ).group_by(Feedback.feedback_type).all()
    feedback_type_dist = {ft: count for ft, count in type_dist}

    # Feedback category distribution
    cat_dist = db.session.query(
        Feedback.category, func.count(Feedback.id)
    ).group_by(Feedback.category).all()
    feedback_category_dist = {cat: count for cat, count in cat_dist}

    # Feedback status distribution
    status_dist = db.session.query(
        Feedback.status, func.count(Feedback.id)
    ).group_by(Feedback.status).all()
    feedback_status_dist = {st: count for st, count in status_dist}
    pending_feedback_count = feedback_status_dist.get('pending', 0)
    reviewing_feedback_count = feedback_status_dist.get('reviewing', 0)
    resolved_feedback_count = feedback_status_dist.get('resolved', 0)
    open_feedback_count = pending_feedback_count + reviewing_feedback_count

    # Rating distribution
    rating_dist = db.session.query(
        Feedback.rating, func.count(Feedback.id)
    ).filter(Feedback.rating.isnot(None)).group_by(Feedback.rating).all()
    feedback_rating_dist = {str(r): count for r, count in rating_dist}

    # Contract type distribution (for admin pie chart)
    ct_dist = db.session.query(
        Analysis.contract_type, func.count(Analysis.id)
    ).filter(Analysis.contract_type.isnot(None)).group_by(Analysis.contract_type).all()
    contract_type_dist = {ct: count for ct, count in ct_dist}

    # User registration trend (7-day)
    daily_registrations = []
    for i in range(6, -1, -1):
        d = date.today() - timedelta(days=i)
        count = User.query.filter(
            func.date(User.created_at) == d,
            User.role != 'anonymous',
        ).count()
        daily_registrations.append({'date': d.isoformat(), 'count': count})

    # Score distribution (for histogram)
    score_dist = db.session.query(
        Analysis.score, func.count(Analysis.id)
    ).filter(Analysis.score.isnot(None)).group_by(Analysis.score).all()
    score_distribution = {str(s): count for s, count in score_dist}

    # Active users (users who analyzed in last 7 days)
    week_ago = date.today() - timedelta(days=7)
    active_users = db.session.query(func.count(func.distinct(Analysis.user_id))).join(
        User, User.id == Analysis.user_id
    ).filter(
        func.date(Analysis.created_at) >= week_ago,
        User.role != 'anonymous',
    ).scalar()

    # Mode distribution
    mode_dist = db.session.query(
        Analysis.analysis_mode, func.count(Analysis.id)
    ).group_by(Analysis.analysis_mode).all()
    mode_distribution = {m: count for m, count in mode_dist}

    # User structure
    role_dist = db.session.query(
        User.role, func.count(User.id)
    ).filter(User.role != 'anonymous').group_by(User.role).all()
    user_role_dist = {role: count for role, count in role_dist}
    avatar_set_count = User.query.filter(User.role != 'anonymous', User.avatar.isnot(None)).count()
    avatar_set_rate = round((avatar_set_count / total_users * 100), 1) if total_users else 0
    new_users_7d = sum(item['count'] for item in daily_registrations)

    return jsonify({
        'total_users': total_users,
        'total_analyses': total_analyses,
        'today_analyses': today_analyses,
        'risk_distribution': risk_distribution,
        'total_feedback': total_feedback,
        'avg_rating': avg_rating,
        'nps': nps,
        'daily_feedback_trend': daily_feedback,
        'daily_analysis_trend': daily_analysis,
        'feedback_type_dist': feedback_type_dist,
        'feedback_category_dist': feedback_category_dist,
        'feedback_status_dist': feedback_status_dist,
        'feedback_rating_dist': feedback_rating_dist,
        'pending_feedback_count': pending_feedback_count,
        'reviewing_feedback_count': reviewing_feedback_count,
        'resolved_feedback_count': resolved_feedback_count,
        'open_feedback_count': open_feedback_count,
        'contract_type_dist': contract_type_dist,
        'daily_registrations': daily_registrations,
        'new_users_7d': new_users_7d,
        'score_distribution': score_distribution,
        'active_users_7d': active_users or 0,
        'mode_distribution': mode_distribution,
        'user_role_dist': user_role_dist,
        'avatar_set_count': avatar_set_count,
        'avatar_set_rate': avatar_set_rate,
        'total_notifications': total_notifications,
        'release_notifications': Notification.query.filter_by(notif_type='release').count(),
        'token_usage': {
            'prompt_tokens': int(token_totals[0] or 0),
            'completion_tokens': int(token_totals[1] or 0),
            'total_tokens': int(token_totals[2] or 0),
            'anonymous_total_tokens': int(anonymous_token_total or 0),
        },
    })


@api_bp.route('/api/admin/users', methods=['GET'])
@get_current_user()
@admin_required
def admin_users(current_user):
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    per_page = min(max(per_page, 1), 100)

    query = db.session.query(
        User, func.count(Analysis.id).label('analysis_count')
    ).outerjoin(Analysis, User.id == Analysis.user_id).filter(
        User.role != 'anonymous'
    ).group_by(User.id).order_by(User.created_at.desc())

    pagination = query.paginate(page=page, per_page=per_page, error_out=False)

    items = []
    for user, count in pagination.items:
        user_dict = user.to_dict()
        user_dict['analysis_count'] = count
        user_dict['referral_count'] = User.query.filter_by(referred_by_user_id=user.id).count()
        items.append(user_dict)

    return jsonify({
        'items': items,
        'total': pagination.total,
        'page': pagination.page,
        'per_page': pagination.per_page,
        'pages': pagination.pages,
    })


@api_bp.route('/api/admin/analyses', methods=['GET'])
@get_current_user()
@admin_required
def admin_analyses(current_user):
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    per_page = min(max(per_page, 1), 100)

    query = Analysis.query.order_by(Analysis.created_at.desc())
    pagination = query.paginate(page=page, per_page=per_page, error_out=False)

    items = [a.to_dict() for a in pagination.items]
    for item in items:
        item['display_user'] = '匿名' if item.get('is_anonymous') else item.get('user_id')

    return jsonify({
        'items': items,
        'total': pagination.total,
        'page': pagination.page,
        'per_page': pagination.per_page,
        'pages': pagination.pages,
    })


# ---------------------------------------------------------------------------
# Feedback Routes
# ---------------------------------------------------------------------------

@api_bp.route('/api/feedback', methods=['POST'])
@get_current_user()
def create_feedback(current_user):
    data = request.get_json(silent=True) or {}
    title = (data.get('title') or '').strip()
    content = (data.get('content') or '').strip()

    if not title or not content:
        return jsonify({'error': 'Title and content are required'}), 400

    feedback = Feedback(
        user_id=current_user.id,
        username=current_user.username,
        title=title,
        content=content,
        feedback_type=(data.get('feedback_type') or 'bug')[:20],
        category=(data.get('category') or '')[:50],
        rating=data.get('rating'),
        contact_email=(data.get('contact_email') or '')[:120],
    )

    try:
        db.session.add(feedback)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Save failed: {str(e)}'}), 500

    return jsonify({'message': 'Feedback submitted', 'feedback': feedback.to_dict()}), 201


@api_bp.route('/api/feedback', methods=['GET'])
@get_current_user()
@admin_required
def list_feedback(current_user):
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    status_filter = request.args.get('status', '').strip()
    type_filter = request.args.get('feedback_type', '').strip()

    query = Feedback.query.order_by(Feedback.created_at.desc())

    if status_filter and status_filter != 'all':
        query = query.filter(Feedback.status == status_filter)
    if type_filter and type_filter != 'all':
        query = query.filter(Feedback.feedback_type == type_filter)

    pagination = query.paginate(page=page, per_page=per_page, error_out=False)

    return jsonify({
        'items': [fb.to_dict() for fb in pagination.items],
        'total': pagination.total,
        'page': pagination.page,
        'per_page': pagination.per_page,
        'pages': pagination.pages,
    })


@api_bp.route('/api/feedback/<int:fb_id>', methods=['GET'])
@get_current_user()
@admin_required
def get_feedback(current_user, fb_id):
    fb = db.session.get(Feedback, fb_id)
    if not fb:
        return jsonify({'error': 'Feedback not found'}), 404
    return jsonify({'feedback': fb.to_dict()})


@api_bp.route('/api/feedback/<int:fb_id>', methods=['PUT'])
@get_current_user()
@admin_required
def update_feedback(current_user, fb_id):
    fb = db.session.get(Feedback, fb_id)
    if not fb:
        return jsonify({'error': 'Feedback not found'}), 404

    data = request.get_json(silent=True) or {}

    if 'status' in data:
        fb.status = data['status'][:20]
    if 'admin_reply' in data:
        fb.admin_reply = data.get('admin_reply', '')

    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Update failed: {str(e)}'}), 500

    return jsonify({'message': 'Updated', 'feedback': fb.to_dict()})


@api_bp.route('/api/feedback/<int:fb_id>', methods=['DELETE'])
@get_current_user()
@admin_required
def delete_feedback(current_user, fb_id):
    fb = db.session.get(Feedback, fb_id)
    if not fb:
        return jsonify({'error': 'Feedback not found'}), 404

    try:
        db.session.delete(fb)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Delete failed: {str(e)}'}), 500

    return jsonify({'message': 'Deleted'})


@api_bp.route('/api/admin/feedback', methods=['GET'])
@get_current_user()
@admin_required
def admin_feedback_list(current_user):
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    feedback_type = request.args.get('type', '').strip()
    status = request.args.get('status', '').strip()
    per_page = min(max(per_page, 1, 100), 100)

    query = Feedback.query.order_by(Feedback.created_at.desc())
    if feedback_type:
        query = query.filter_by(feedback_type=feedback_type)
    if status:
        query = query.filter_by(status=status)

    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    items = [f.to_dict() for f in pagination.items]

    return jsonify({
        'items': items,
        'total': pagination.total,
        'page': pagination.page,
        'per_page': pagination.per_page,
        'pages': pagination.pages,
    })


@api_bp.route('/api/quota', methods=['GET'])
@get_current_user()
def get_user_quota(current_user):
    """Get current user's remaining analysis quota for today."""
    quota = UserQuota.get_today_quota(current_user.id)
    user = db.session.get(User, current_user.id)
    analysis_limit = UserQuota.get_effective_limit(user, 'analysis')
    compare_limit = UserQuota.get_effective_limit(user, 'compare')
    followup_limit = UserQuota.get_effective_limit(user, 'followup')
    analysis_used = UserQuota.get_usage(user, 'analysis')
    compare_used = UserQuota.get_usage(user, 'compare')
    followup_used = UserQuota.get_usage(user, 'followup')
    analysis_remaining = max(0, analysis_limit - analysis_used)
    compare_remaining = max(0, compare_limit - compare_used)
    followup_remaining = max(0, followup_limit - followup_used)
    return jsonify({
        'analysis': {
            'used': analysis_used,
            'remaining': analysis_remaining,
            'limit': analysis_limit,
        },
        'compare': {
            'used': compare_used,
            'remaining': compare_remaining,
            'limit': compare_limit,
        },
        'followup': {
            'used': followup_used,
            'remaining': followup_remaining,
            'limit': followup_limit,
        },
        'analysis_used': analysis_used,
        'analysis_remaining': analysis_remaining,
        'analysis_limit': analysis_limit,
        'compare_used': compare_used,
        'compare_remaining': compare_remaining,
        'compare_limit': compare_limit,
        'followup_used': followup_used,
        'followup_remaining': followup_remaining,
        'followup_limit': followup_limit,
        'bonus_credits': quota.bonus_credits or 0,
        'analysis_bonus_credits': user.reward_analysis_credits or 0,
        'compare_bonus_credits': user.reward_compare_credits or 0,
        'followup_bonus_credits': user.reward_followup_credits or 0,
        'reward_analysis_credits': user.reward_analysis_credits or 0,
        'reward_compare_credits': user.reward_compare_credits or 0,
        'reward_followup_credits': user.reward_followup_credits or 0,
        'bonus_granted_date': quota.bonus_granted_date.isoformat() if quota.bonus_granted_date else None,
    })


# ---------------------------------------------------------------------------
# Notification API
# ---------------------------------------------------------------------------

def _ensure_user_notif_state(notification_id, user_id):
    """确保每个用户对每条通知都有状态记录（懒创建）。"""
    state = UserNotification.query.filter_by(
        notification_id=notification_id, user_id=user_id
    ).first()
    if not state:
        state = UserNotification(
            notification_id=notification_id,
            user_id=user_id,
            is_read=False,
        )
        db.session.add(state)
        db.session.commit()
    return state


def _format_time_ago(dt, lang='zh'):
    """格式化相对时间。"""
    if not dt:
        return ''
    now = datetime.utcnow()
    diff = now - dt
    seconds = int(diff.total_seconds())
    if lang == 'en':
        if seconds < 60:
            return 'Just now'
        if seconds < 3600:
            return f'{seconds // 60}m ago'
        if seconds < 86400:
            return f'{seconds // 3600}h ago'
        if seconds < 2592000:
            return f'{seconds // 86400}d ago'
        return dt.strftime('%Y-%m-%d')
    else:
        if seconds < 60:
            return '刚刚'
        if seconds < 3600:
            return f'{seconds // 60} 分钟前'
        if seconds < 86400:
            return f'{seconds // 3600} 小时前'
        if seconds < 172800:
            return '昨天'
        if seconds < 2592000:
            return f'{seconds // 86400} 天前'
        return dt.strftime('%Y-%m-%d')


def _localize_notification(item, lang='zh'):
    if lang != 'en':
        return item
    title = item.get('title') or ''
    if title.startswith('版本更新：'):
        version = title.split('：', 1)[1].strip()
        item = dict(item)
        item.update({
            'title': f'Release Update: {version}' if version else 'Release Update',
            'summary': 'This release includes a comparison page fix, version-wide notifications, and daily quota controls.',
        })
        return item
    translations = {
        '欢迎使用 DocAI': {
            'title': 'Welcome to DocAI',
            'summary': 'Your DocAI smart contract analysis workspace is ready. Start your first contract analysis.',
        },
        '新功能上线：合同对比': {
            'title': 'New Feature: Contract Comparison',
            'summary': 'Contract comparison now supports smart two-document review. Try difference detection and risk assessment.',
        },
        '系统维护通知': {
            'title': 'Scheduled Maintenance',
            'summary': 'The system will undergo routine maintenance this Sunday from 2:00 to 4:00 AM. Service may be briefly interrupted.',
        },
    }
    mapped = translations.get(item.get('title'))
    if mapped:
        item = dict(item)
        item.update(mapped)
    return item


@api_bp.route('/api/notifications', methods=['GET'])
@get_current_user()
def list_notifications(current_user):
    """获取当前用户的通知列表（含已读状态），支持 tab 筛选。"""
    tab = request.args.get('tab', 'all')
    lang = request.args.get('lang', 'zh')

    # 查询所有未被当前用户软删除的通知
    query = Notification.query.order_by(Notification.created_at.desc())
    notifications = query.all()

    result = []
    for n in notifications:
        state = _ensure_user_notif_state(n.id, current_user.id)
        if state.deleted:
            continue
        item = n.to_dict(current_user.id)
        item['time_ago'] = _format_time_ago(n.created_at, lang)
        # tab 筛选
        if tab == 'unread' and item['read']:
            continue
        if tab == 'announce' and n.notif_type not in ('system', 'release'):
            continue
        if tab == 'alert' and n.notif_type != 'alert':
            continue
        result.append(_localize_notification(item, lang))

    unread_count = UserNotification.query.filter_by(
        user_id=current_user.id, is_read=False, deleted=False
    ).count()
    # 对没有状态记录的通知，也算未读
    total_unread = 0
    for n in notifications:
        state = UserNotification.query.filter_by(
            notification_id=n.id, user_id=current_user.id
        ).first()
        if not state or (not state.is_read and not state.deleted):
            total_unread += 1

    return jsonify({
        'notifications': result,
        'unread_count': total_unread,
        'total': len(result),
    })


@api_bp.route('/api/notifications/unread-count', methods=['GET'])
@get_current_user()
def unread_count(current_user):
    """获取未读通知数量（用于红点展示）。"""
    notifications = Notification.query.all()
    count = 0
    for n in notifications:
        state = UserNotification.query.filter_by(
            notification_id=n.id, user_id=current_user.id
        ).first()
        if not state or (not state.is_read and not state.deleted):
            count += 1
    return jsonify({'unread_count': count})


@api_bp.route('/api/notifications/<int:notif_id>/read', methods=['POST'])
@get_current_user()
def mark_notification_read(current_user, notif_id):
    """标记单条通知为已读。"""
    state = _ensure_user_notif_state(notif_id, current_user.id)
    state.is_read = True
    state.read_at = datetime.utcnow()
    db.session.commit()
    return jsonify({'message': 'ok', 'read': True})


@api_bp.route('/api/notifications/read-all', methods=['POST'])
@get_current_user()
def mark_all_notifications_read(current_user):
    """标记所有通知为已读。"""
    notifications = Notification.query.all()
    for n in notifications:
        state = _ensure_user_notif_state(n.id, current_user.id)
        if not state.is_read:
            state.is_read = True
            state.read_at = datetime.utcnow()
    db.session.commit()
    return jsonify({'message': 'ok'})


@api_bp.route('/api/notifications/<int:notif_id>', methods=['DELETE'])
@get_current_user()
def delete_notification(current_user, notif_id):
    """软删除：当前用户隐藏该通知。"""
    state = _ensure_user_notif_state(notif_id, current_user.id)
    state.deleted = True
    db.session.commit()
    return jsonify({'message': 'ok'})


@api_bp.route('/api/notifications/clear-all', methods=['DELETE'])
@get_current_user()
def clear_all_notifications(current_user):
    """清空当前用户的所有通知（软删除）。"""
    notifications = Notification.query.all()
    for n in notifications:
        state = _ensure_user_notif_state(n.id, current_user.id)
        state.deleted = True
    db.session.commit()
    return jsonify({'message': 'ok'})


@api_bp.route('/api/admin/notifications/broadcast', methods=['POST'])
@get_current_user()
def broadcast_notification(current_user):
    """管理员广播通知：向所有用户发送一条通知。"""
    if current_user.role != 'admin':
        return jsonify({'error': '权限不足'}), 403
    data = request.get_json() or {}
    title = (data.get('title') or '').strip()
    if not title:
        return jsonify({'error': '标题不能为空'}), 400
    summary = (data.get('summary') or '').strip()
    notif_type = data.get('notif_type', 'system')
    icon = data.get('icon')
    link = data.get('link')

    notif = Notification(
        title=title,
        summary=summary,
        notif_type=notif_type,
        icon=icon,
        link=link,
    )
    db.session.add(notif)
    db.session.commit()
    # 所有用户的通知状态会在首次查询时懒创建
    return jsonify({
        'message': '通知已广播给所有用户',
        'notification': notif.to_dict(),
    })


@api_bp.route('/api/admin/notifications', methods=['GET'])
@get_current_user()
def admin_notifications(current_user):
    """管理员查看系统通知与同步状态。"""
    if current_user.role != 'admin':
        return jsonify({'error': '权限不足'}), 403

    total_users = User.query.count()
    items = []
    for n in Notification.query.order_by(Notification.created_at.desc()).limit(30).all():
        read_count = UserNotification.query.filter_by(notification_id=n.id, is_read=True, deleted=False).count()
        delivered_count = UserNotification.query.filter_by(notification_id=n.id, deleted=False).count()
        unread_count = max(0, total_users - read_count)
        items.append({
            **n.to_dict(),
            'read_count': read_count,
            'delivered_count': delivered_count,
            'unread_count': unread_count,
            'delivery_rate': round((delivered_count / total_users * 100), 1) if total_users else 0,
        })

    return jsonify({
        'items': items,
        'total_users': total_users,
        'total_notifications': len(items),
    })


# ---------------------------------------------------------------------------
# User Dashboard API
# ---------------------------------------------------------------------------

@api_bp.route('/api/dashboard', methods=['GET'])
@get_current_user()
def user_dashboard(current_user):
    """Get data for the user's personal dashboard."""
    from models import UserQuota

    # Quota
    quota = UserQuota.get_today_quota(current_user.id)
    analysis_limit = UserQuota.get_effective_limit(current_user, 'analysis')
    compare_limit = UserQuota.get_effective_limit(current_user, 'compare')
    followup_limit = UserQuota.get_effective_limit(current_user, 'followup')
    analysis_used = UserQuota.get_usage(current_user, 'analysis')
    compare_used = UserQuota.get_usage(current_user, 'compare')
    followup_used = UserQuota.get_usage(current_user, 'followup')
    analysis_remaining = max(0, analysis_limit - analysis_used)
    compare_remaining = max(0, compare_limit - compare_used)
    followup_remaining = max(0, followup_limit - followup_used)

    # Recent analyses (last 5)
    recent = Analysis.query.filter_by(user_id=current_user.id).order_by(
        Analysis.created_at.desc()
    ).limit(5).all()
    recent_items = [a.to_dict() for a in recent]

    # Total stats
    total_analyses = Analysis.query.filter_by(user_id=current_user.id).count()
    total_compares = ContractCompare.query.filter_by(user_id=current_user.id).count()

    # 7-day trend
    daily_trend = []
    for i in range(6, -1, -1):
        d = date.today() - timedelta(days=i)
        count = Analysis.query.filter(
            Analysis.user_id == current_user.id,
            func.date(Analysis.created_at) == d
        ).count()
        daily_trend.append({'date': d.isoformat(), 'count': count})

    # Risk distribution of user's analyses
    risk_dist_query = db.session.query(
        Analysis.risk_level, func.count(Analysis.id)
    ).filter(Analysis.user_id == current_user.id).group_by(Analysis.risk_level).all()
    risk_dist = {level: count for level, count in risk_dist_query}

    # Contract type distribution
    type_dist_query = db.session.query(
        Analysis.contract_type, func.count(Analysis.id)
    ).filter(
        Analysis.user_id == current_user.id,
        Analysis.contract_type.isnot(None)
    ).group_by(Analysis.contract_type).all()
    type_dist = {t: count for t, count in type_dist_query}

    # Average score
    avg_score = db.session.query(func.avg(Analysis.score)).filter(
        Analysis.user_id == current_user.id,
        Analysis.score.isnot(None)
    ).scalar()
    avg_score = round(float(avg_score), 1) if avg_score else None

    # High risk count (needs attention)
    high_risk_count = Analysis.query.filter(
        Analysis.user_id == current_user.id,
        Analysis.risk_level == 'high'
    ).count()

    referral_count = User.query.filter_by(referred_by_user_id=current_user.id).count()
    invite_code = current_user.invite_code
    if not invite_code:
        invite_prefix = current_app.config.get('REFERRAL_CODE_PREFIX', 'DCAI')
        invite_code = f'{invite_prefix}{current_user.id:06d}'
        current_user.invite_code = invite_code
        db.session.commit()
    referral_url = f"{request.url_root.rstrip('/')}/login?ref={invite_code}"

    return jsonify({
        'quota': {
            'analysis_used': analysis_used,
            'analysis_remaining': analysis_remaining,
            'analysis_limit': analysis_limit,
            'compare_used': compare_used,
            'compare_remaining': compare_remaining,
            'compare_limit': compare_limit,
            'followup_used': followup_used,
            'followup_remaining': followup_remaining,
            'followup_limit': followup_limit,
            'bonus_credits': quota.bonus_credits or 0,
            'analysis_bonus_credits': current_user.reward_analysis_credits or 0,
            'compare_bonus_credits': current_user.reward_compare_credits or 0,
            'followup_bonus_credits': current_user.reward_followup_credits or 0,
            'reward_analysis_credits': current_user.reward_analysis_credits or 0,
            'reward_compare_credits': current_user.reward_compare_credits or 0,
            'reward_followup_credits': current_user.reward_followup_credits or 0,
        },
        'quota_used': analysis_used,
        'quota_remaining': analysis_remaining,
        'quota_limit': analysis_limit,
        'recent_analyses': recent_items,
        'referral': {
            'invite_code': invite_code,
            'invite_url': referral_url,
            'referral_count': referral_count,
            'referral_bonus': current_app.config.get('REFERRAL_BONUS_CREDITS', 2),
        },
        'stats': {
            'total_analyses': total_analyses,
            'total_compares': total_compares,
            'avg_score': avg_score,
            'high_risk_count': high_risk_count,
        },
        'daily_trend': daily_trend,
        'risk_distribution': risk_dist,
        'contract_type_distribution': type_dist,
    })


@api_bp.route('/api/activity', methods=['GET'])
@get_current_user()
def user_activity(current_user):
    """Return real account activity for the current user."""
    items = []

    if current_user.created_at:
        items.append({
            'type': 'register',
            'status': 'success',
            'action_key': 'dash.activity_act_register',
            'meta': current_user.email,
            'created_at': current_user.created_at.isoformat(),
        })

    analyses = Analysis.query.filter_by(user_id=current_user.id).order_by(
        Analysis.created_at.desc()
    ).limit(30).all()
    for analysis in analyses:
        items.append({
            'type': 'analysis',
            'status': 'success',
            'action_key': 'dash.activity_act_analysis',
            'meta': analysis.filename or ('Pasted text' if (analysis.language or 'zh') == 'en' else '粘贴文本'),
            'created_at': analysis.created_at.isoformat() if analysis.created_at else None,
        })

    compares = ContractCompare.query.filter_by(user_id=current_user.id).order_by(
        ContractCompare.created_at.desc()
    ).limit(30).all()
    for compare in compares:
        items.append({
            'type': 'compare',
            'status': 'success',
            'action_key': 'dash.activity_act_compare',
            'meta': 'Contract comparison',
            'created_at': compare.created_at.isoformat() if compare.created_at else None,
        })

    items.sort(key=lambda item: item.get('created_at') or '', reverse=True)
    return jsonify({'items': items[:50], 'total': len(items)})


# ---------------------------------------------------------------------------
# Report Export API
# ---------------------------------------------------------------------------

@api_bp.route('/api/analysis/<int:analysis_id>/export', methods=['GET'])
@get_current_user()
def export_analysis_pdf(current_user, analysis_id):
    """Export an analysis as a printable HTML page or downloadable DOCX file."""
    analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first()
    if not analysis:
        return jsonify({'error': '分析记录不存在'}), 404

    # Parse result
    result = {}
    if analysis.result:
        try:
            result = json.loads(analysis.result)
        except (json.JSONDecodeError, TypeError):
            pass

    risk_items = []
    if analysis.risk_items:
        try:
            risk_items = json.loads(analysis.risk_items)
        except (json.JSONDecodeError, TypeError):
            pass

    suggestions = []
    if analysis.suggestions:
        try:
            suggestions = json.loads(analysis.suggestions)
        except (json.JSONDecodeError, TypeError):
            pass

    contract_type_map = {
        'labor': '劳动合同', 'rental': '房屋租赁', 'purchase': '买卖合同',
        'service': '服务合同', 'nda': '保密协议', 'loan': '借款合同',
        'partnership': '合伙协议', 'franchise': '加盟合同', 'agency': '委托合同',
        'construction': '建设工程', 'insurance': '保险合同', 'other': '其他',
    }
    contract_type_label = contract_type_map.get(analysis.contract_type, '其他')

    risk_colors = {'high': '#D45050', 'medium': '#E8A33D', 'low': '#2D9D78'}
    risk_labels = {'high': '高风险', 'medium': '中风险', 'low': '低风险'}
    mode_labels = {'risk': '风险分析', 'summary': '合同摘要', 'plain': '通俗解读'}

    if request.args.get('format', '').lower() == 'docx':
        return _export_analysis_docx(
            analysis, contract_type_label, risk_labels, mode_labels,
            risk_items, suggestions,
        )

    mode_label = mode_labels.get(analysis.analysis_mode, '合同分析')

    # Build printable HTML
    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>DocAI 合同分析报告 - {escape_html(analysis.filename or "合同分析")}</title>
<style>
  @page {{ size: A4; margin: 20mm; }}
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: 'Inter', 'Noto Sans SC', -apple-system, sans-serif; color: #1A2332; line-height: 1.6; padding: 40px; max-width: 800px; margin: 0 auto; }}
  .header {{ text-align: center; margin-bottom: 32px; padding-bottom: 20px; border-bottom: 2px solid #1B365D; }}
  .header h1 {{ font-size: 24px; color: #1B365D; margin-bottom: 8px; }}
  .header p {{ font-size: 13px; color: #8892A0; }}
  .meta {{ display: flex; gap: 24px; margin-bottom: 24px; flex-wrap: wrap; }}
  .meta-item {{ font-size: 13px; color: #5A6677; }}
  .meta-item strong {{ color: #1A2332; }}
  .score-section {{ text-align: center; padding: 24px; background: #F1F4F8; border-radius: 12px; margin-bottom: 24px; }}
  .score-number {{ font-size: 48px; font-weight: 700; color: {risk_colors.get(analysis.risk_level, "#2D9D78")}; }}
  .score-label {{ font-size: 14px; color: #5A6677; }}
  .summary-box {{ background: #F1F4F8; border-left: 4px solid #2D9D78; padding: 16px; border-radius: 8px; margin-bottom: 24px; font-size: 14px; color: #1A2332; }}
  .section-title {{ font-size: 18px; font-weight: 700; color: #1B365D; margin: 24px 0 16px; padding-bottom: 8px; border-bottom: 1px solid #E8ECF1; }}
  .risk-item {{ border: 1px solid #E8ECF1; border-radius: 8px; padding: 16px; margin-bottom: 12px; page-break-inside: avoid; }}
  .risk-item-header {{ display: flex; align-items: center; gap: 8px; margin-bottom: 8px; }}
  .risk-badge {{ font-size: 11px; font-weight: 600; padding: 2px 8px; border-radius: 4px; color: white; }}
  .risk-name {{ font-weight: 600; font-size: 14px; }}
  .risk-field {{ font-size: 13px; color: #5A6677; margin-bottom: 4px; }}
  .risk-field strong {{ color: #1A2332; }}
  .suggestion-item {{ background: #E6F5F0; padding: 10px 14px; border-radius: 8px; margin-bottom: 8px; font-size: 13px; color: #1F7D5C; }}
  .footer {{ margin-top: 40px; padding-top: 16px; border-top: 1px solid #E8ECF1; text-align: center; font-size: 11px; color: #8892A0; }}
  @media print {{ body {{ padding: 0; }} .no-print {{ display: none; }} }}
</style>
</head>
<body>
<div class="header">
  <h1>DocAI 合同分析报告</h1>
  <p>AI 智能合同分析平台 | 生成时间: {datetime.utcnow().strftime("%Y-%m-%d %H:%M")}</p>
</div>
<div class="meta">
  <div class="meta-item"><strong>文件名:</strong> {escape_html(analysis.filename or "粘贴文本")}</div>
  <div class="meta-item"><strong>分析模式:</strong> {mode_label}</div>
  <div class="meta-item"><strong>合同类型:</strong> {contract_type_label}</div>
  <div class="meta-item"><strong>分析时间:</strong> {analysis.created_at.strftime("%Y-%m-%d %H:%M") if analysis.created_at else "-"}</div>
</div>'''

    # Score section
    score_val = analysis.score if analysis.score is not None else '--'
    risk_label = risk_labels.get(analysis.risk_level, '未知')
    html += f'''
<div class="score-section">
  <div class="score-number">{score_val}</div>
  <div class="score-label">综合评分 · {risk_label}</div>
</div>'''

    # One-line summary
    if analysis.one_line_summary:
        html += f'''
<div class="summary-box">
  <strong>核心摘要:</strong> {escape_html(analysis.one_line_summary)}
</div>'''

    # Risk items
    if risk_items:
        html += '<div class="section-title">风险条款详情</div>'
        for idx, item in enumerate(risk_items, 1):
            if not isinstance(item, dict):
                continue
            rl = item.get('risk_level', 'medium')
            color = risk_colors.get(rl, '#8892A0')
            rl_label = risk_labels.get(rl, '未知')
            name = escape_html(item.get('clause_name', f'风险条款 {idx}'))
            location = escape_html(item.get('clause_location', ''))
            desc = escape_html(item.get('description', ''))
            plain = escape_html(item.get('plain_explanation', ''))
            suggestion = escape_html(item.get('suggestion', ''))
            legal = escape_html(item.get('legal_basis', ''))

            html += f'''
<div class="risk-item">
  <div class="risk-item-header">
    <span class="risk-badge" style="background: {color};">{rl_label}</span>
    <span class="risk-name">{name}</span>
    {f'<span style="font-size:12px;color:#8892A0;">{location}</span>' if location else ''}
  </div>
  <div class="risk-field">{desc}</div>'''
            if plain:
                html += f'''<div class="risk-field" style="color:#2D9D78;"><strong>通俗解读:</strong> {plain}</div>'''
            if suggestion:
                html += f'''<div class="suggestion-item"><strong>修改建议:</strong> {suggestion}</div>'''
            if legal:
                html += f'''<div class="risk-field"><strong>法律依据:</strong> {legal}</div>'''
            html += '</div>'

    # Suggestions
    if suggestions:
        html += '<div class="section-title">总体改进建议</div>'
        for s in suggestions:
            html += f'<div class="suggestion-item">{escape_html(str(s))}</div>'

    # Footer
    html += '''
<div class="footer">
  <p>本报告由 DocAI AI 智能合同分析平台自动生成，仅供参考，不构成法律意见。</p>
  <p>DocAI &copy; 2026 | 用 AI 守护每一份合同权益</p>
</div>
</body></html>'''

    return html, 200, {'Content-Type': 'text/html; charset=utf-8'}


def _export_analysis_docx(analysis, contract_type_label, risk_labels, mode_labels, risk_items, suggestions):
    """Create a Word report using the same data shown in the printable report."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError:
        current_app.logger.error('python-docx is not installed; DOCX export is unavailable.')
        return jsonify({'error': 'Word 导出组件暂不可用，请稍后重试'}), 503

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    normal_style = document.styles['Normal']
    normal_style.font.name = 'Microsoft YaHei'
    normal_style.font.size = Pt(10.5)

    title = document.add_heading('DocAI 合同分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(
        f'AI 智能合同分析平台 | 生成时间：{datetime.utcnow().strftime("%Y-%m-%d %H:%M")} UTC'
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading('报告概览', level=1)
    metadata = document.add_table(rows=0, cols=2)
    metadata.style = 'Table Grid'
    metadata_rows = (
        ('文件名', analysis.filename or '粘贴文本'),
        ('分析模式', mode_labels.get(analysis.analysis_mode, '合同分析')),
        ('合同类型', contract_type_label),
        ('分析时间', analysis.created_at.strftime('%Y-%m-%d %H:%M') if analysis.created_at else '-'),
        ('综合评分', str(analysis.score) if analysis.score is not None else '--'),
        ('风险等级', risk_labels.get(analysis.risk_level, '未知')),
    )
    for label, value in metadata_rows:
        cells = metadata.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)

    if analysis.one_line_summary:
        document.add_heading('核心摘要', level=1)
        document.add_paragraph(analysis.one_line_summary)

    if risk_items:
        document.add_heading('风险条款详情', level=1)
        for idx, item in enumerate(risk_items, 1):
            if not isinstance(item, dict):
                continue
            risk_level = item.get('risk_level', 'medium')
            name = item.get('clause_name') or f'风险条款 {idx}'
            document.add_heading(f'{idx}. {name}（{risk_labels.get(risk_level, "未知")}）', level=2)
            _add_docx_report_field(document, '条款位置', item.get('clause_location'))
            _add_docx_report_field(document, '风险说明', item.get('description'))
            _add_docx_report_field(document, '通俗解读', item.get('plain_explanation'))
            _add_docx_report_field(document, '修改建议', item.get('suggestion'))
            _add_docx_report_field(document, '法律依据', item.get('legal_basis'))

    if suggestions:
        document.add_heading('总体改进建议', level=1)
        for suggestion in suggestions:
            document.add_paragraph(str(suggestion), style='List Bullet')

    document.add_paragraph('本报告由 DocAI AI 智能合同分析平台自动生成，仅供参考，不构成法律意见。')

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return send_file(
        output,
        as_attachment=True,
        download_name=f'DocAI-Contract-Analysis-{analysis.id}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


def _add_docx_report_field(document, label, value):
    if value:
        paragraph = document.add_paragraph()
        paragraph.add_run(f'{label}：').bold = True
        paragraph.add_run(str(value))


def escape_html(text):
    """Basic HTML escaping for PDF export."""
    if not text:
        return ''
    return str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')


# ---------------------------------------------------------------------------
# Page routes (render Jinja2 templates)
# ---------------------------------------------------------------------------

@api_bp.route('/')
def page_index():
    return render_template('index.html', active_nav='home')

@api_bp.route('/dashboard')
def page_dashboard():
    # 未登录用户访问个人主页时，显示登录引导页
    from flask import request
    token = request.cookies.get('token')
    if not token:
        return render_template('login_required.html', active_nav='dashboard')
    try:
        from auth import decode_token
        user = decode_token(token)
    except Exception:
        return render_template('login_required.html', active_nav='dashboard')
    return render_template('dashboard.html', active_nav='dashboard')

@api_bp.route('/analyze')
def page_analyze():
    return render_template('analyze.html', active_nav='analyze')

@api_bp.route('/archive')
def page_archive():
    return render_template('archive.html', active_nav='archive')

@api_bp.route('/compare')
def page_compare():
    return render_template('compare.html', active_nav='compare')

@api_bp.route('/pricing')
def page_pricing():
    return render_template('pricing.html', active_nav='pricing')

@api_bp.route('/about')
def page_about():
    return render_template('about.html', active_nav='about')

@api_bp.route('/login')
def page_login():
    return render_template('login.html')

@api_bp.route('/admin')
def page_admin():
    return render_template('admin.html', active_nav='admin')

@api_bp.route('/detail/<int:analysis_id>')
def page_detail(analysis_id):
    return render_template('detail.html', analysis_id=analysis_id)
